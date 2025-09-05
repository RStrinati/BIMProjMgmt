from typing import Optional
import os

try:
    from docx import Document  # python-docx
except Exception:
    Document = None

from .dao import list_project_document_sections, list_project_clauses


def render_docx(conn, project_document_id: int) -> Optional[str]:
    """Render a simple DOCX concatenating sections and clauses. Returns file path or None."""
    # Fallback to txt if python-docx not available
    sections = list_project_document_sections(conn, project_document_id)
    clauses = list_project_clauses(conn, project_document_id)
    clauses_by_section = {}
    for c in clauses:
        clauses_by_section.setdefault(c.get('section_code'), []).append(c)

    out_dir = os.path.join(os.getcwd(), 'docs')
    try:
        os.makedirs(out_dir, exist_ok=True)
    except Exception:
        pass

    if Document is None:
        # Write a TXT file as fallback
        path = os.path.join(out_dir, f"project_doc_{project_document_id}.txt")
        with open(path, 'w', encoding='utf-8') as f:
            for s in sections:
                f.write(f"{s.get('code') or ''} {s.get('title')}\n")
                for c in [x for x in clauses if x['section_code'] == s.get('code')]:
                    title = c.get('title') or ''
                    body = c.get('body_md') or ''
                    f.write(f"  - {c.get('code') or ''} {title}\n")
                    if body:
                        f.write(f"{body}\n\n")
        return path

    # Build DOCX
    doc = Document()
    doc.add_heading('Project Document', 0)
    for s in sections:
        doc.add_heading(f"{s.get('code') or ''} {s.get('title')}", level=1)
        for c in [x for x in clauses if x['section_code'] == s.get('code')]:
            if c.get('title'):
                doc.add_heading(f"{c.get('code') or ''} {c.get('title')}", level=2)
            body = c.get('body_md') or ''
            if body:
                doc.add_paragraph(body)
    path = os.path.join(out_dir, f"project_doc_{project_document_id}.docx")
    doc.save(path)
    return path

